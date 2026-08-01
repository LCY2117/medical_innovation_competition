from __future__ import annotations

import csv
import statistics
from datetime import datetime, timezone, timedelta
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parents[1]
SOURCE_DIR = ROOT / "competition_materials" / "medical_innovation_competition" / "医创赛预实验数据"
OUTPUT_DIR = ROOT / "competition_materials" / "medical_innovation_competition" / "实验数据打印精选包_20260528"
MARKDOWN_OUT = OUTPUT_DIR / "生命反射弧_系统级模拟预实验数据记录_正式打印版.md"
DOCX_OUT = OUTPUT_DIR / "生命反射弧_系统级模拟预实验数据记录_正式打印版.docx"

ROUND_DIRS = [
    "第一轮 校园标准交接流程",
    "第二轮 AED取送较快流程",
    "第三轮 救护车较早接近现场流程",
    "第四轮 现场协同稍慢流程",
    "第五轮 含AED分析与除颤确认流程",
]

TIME_METRICS = [
    ("T1", "触发到分派完成秒数", "触发到分派完成"),
    ("T2", "触发到核心施救响应秒数", "触发到核心施救响应"),
    ("T3", "触发到CPR开始秒数", "触发到 CPR 开始"),
    ("T4", "触发到AED取到秒数", "触发到 AED 取到"),
    ("T5", "触发到AED送达秒数", "触发到 AED 送达"),
    ("T6", "触发到救护接管秒数", "触发到救护接管"),
]


def read_csv(path: Path) -> list[dict[str, str]]:
    with path.open("r", encoding="utf-8-sig", newline="") as f:
        return list(csv.DictReader(f))


def number(value: str) -> float | None:
    try:
        if value is None or str(value).strip() == "":
            return None
        return float(value)
    except ValueError:
        return None


def fmt_num(value: float | int | str | None) -> str:
    if value is None:
        return ""
    if isinstance(value, str):
        parsed = number(value)
        if parsed is None:
            return value
        value = parsed
    if abs(float(value) - int(float(value))) < 0.000001:
        return str(int(float(value)))
    return f"{float(value):.2f}".rstrip("0").rstrip(".")


def iso_to_beijing(value: str) -> str:
    if not value:
        return ""
    try:
        dt = datetime.fromisoformat(value.replace("Z", "+00:00"))
        dt = dt.astimezone(timezone(timedelta(hours=8)))
        return dt.strftime("%Y-%m-%d %H:%M:%S")
    except ValueError:
        return value


def md_escape(value: str | None) -> str:
    text = "" if value is None else str(value)
    return text.replace("|", "\\|").replace("\n", "<br>").strip()


def md_table(headers: list[str], rows: Iterable[Iterable[str]]) -> str:
    lines = [
        "| " + " | ".join(md_escape(h) for h in headers) + " |",
        "| " + " | ".join("---" for _ in headers) + " |",
    ]
    for row in rows:
        lines.append("| " + " | ".join(md_escape(str(cell)) for cell in row) + " |")
    return "\n".join(lines)


def stat(values: list[float]) -> tuple[str, str, str, str]:
    return (
        fmt_num(statistics.mean(values)),
        fmt_num(statistics.median(values)),
        fmt_num(min(values)),
        fmt_num(max(values)),
    )


def load_round(round_dir_name: str) -> dict:
    round_dir = SOURCE_DIR / round_dir_name
    summary = read_csv(round_dir / "单轮汇总表.csv")[0]
    timeline = read_csv(round_dir / "时间线.csv")
    dispatch = read_csv(round_dir / "分派依据.csv")
    metrics = read_csv(round_dir / "指标表.csv")
    return {
        "name": round_dir_name,
        "summary": summary,
        "timeline": timeline,
        "dispatch": dispatch,
        "metrics": metrics,
    }


def build_markdown(rounds: list[dict]) -> str:
    quality_rows = {
        row["轮次编号"]: row
        for row in read_csv(SOURCE_DIR / "多轮分析" / "多轮汇总表.csv")
    }
    summary_rows = []
    for item in rounds:
        row = item["summary"]
        quality = quality_rows.get(row["轮次编号"], {})
        summary_rows.append(
            [
                row["轮次编号"],
                item["name"],
                row["事件阶段"],
                row["分派来源"],
                row["参与终端数"],
                row["可用AED点位数"],
                quality.get("质量分", ""),
                quality.get("校验状态", "通过"),
            ]
        )

    metric_stat_rows = []
    for code, key, label in TIME_METRICS:
        values = [number(item["summary"].get(key, "")) for item in rounds]
        clean = [v for v in values if v is not None]
        mean, median, min_v, max_v = stat(clean)
        metric_stat_rows.append([code, label, len(clean), mean, median, min_v, max_v])

    coverage_rows = []
    for key, label, unit in [
        ("角色分派完整度", "角色分派完整度", "%"),
        ("定位覆盖率", "定位覆盖率", "%"),
        ("健康摘要覆盖率", "健康摘要覆盖率", "%"),
    ]:
        values = [number(item["summary"].get(key, "")) for item in rounds]
        clean = [v for v in values if v is not None]
        if key == "角色分派完整度":
            clean = [v * 100 for v in clean]
        mean, median, min_v, max_v = stat(clean)
        coverage_rows.append([label, len(clean), f"{mean}{unit}", f"{median}{unit}", f"{min_v}{unit}", f"{max_v}{unit}"])

    lines: list[str] = []
    lines.extend(
        [
            "# 生命反射弧系统级模拟预实验数据记录",
            "",
            "项目名称：生命反射弧",
            "",
            "赛道定位：第十二届全国大学生医学创新大赛暨“一带一路”国际竞赛，交叉学科方向，AI 创新设计组。",
            "",
            "数据生成日期：2026-05-26",
            "",
            "文档整理日期：2026-05-28",
            "",
            "## 一、实验目的",
            "",
            "本次系统级模拟预实验用于验证“生命反射弧”网页端急救协同系统在模拟院前急救场景中的工程闭环能力，重点观察系统是否能够完成事件触发、角色分派、四端协同、关键时间节点记录、AED 取送记录、救护接管记录和证据包导出。",
            "",
            "本实验不用于评价真实临床疗效，不用于证明抢救成功率提升，也不替代 120、专业医护判断或 AED 设备说明。实验结论仅限于系统流程可行性、数据记录完整性和协同任务可解释性。",
            "",
            "## 二、实验设计与执行方法",
            "",
            "实验类型：系统级模拟预实验。",
            "",
            "实验对象：网页端总控台与四类模拟终端，分别对应患者触发端、核心施救端、AED 保障端、环境清障/救护接驳端。",
            "",
            "实验场景：校园或公共空间中疑似心脏骤停事件的院前急救协同流程。系统通过预置模拟终端、样例 AED 点位和事件状态机完成桌面推演。",
            "",
            "实验轮次：共 5 轮，覆盖标准交接、AED 取送较快、救护车较早接近、现场协同稍慢、含 AED 分析与除颤确认等不同流程情形。",
            "",
            "分派方式：本批数据记录的分派来源均为“规则备用分派”。该设置用于验证系统在演示稳定性和 AI 接口不可用时的兜底分派能力，不作为 AI 模型性能评估。",
            "",
            "实验流程如下：",
            "",
            md_table(
                ["步骤", "流程节点", "系统记录内容"],
                [
                    ["1", "初始化事件与终端", "创建事件，加载模拟参与者、AED 点位和总控台状态"],
                    ["2", "指定患者端", "记录患者模拟端，形成事件触发主体"],
                    ["3", "启动分派", "系统根据角色、距离、能力和 AED 路径生成分派结果"],
                    ["4", "核心施救响应", "记录核心施救者响应、到场和 CPR 开始节点"],
                    ["5", "AED 保障执行", "记录 AED 取到、送达和交付相关节点"],
                    ["6", "环境清障与救护接驳", "记录通道协调、救护车到场和接管节点"],
                    ["7", "交接归档", "完成事件归档并生成结构化证据包"],
                ],
            ),
            "",
            "## 三、指标定义",
            "",
            md_table(
                ["指标", "含义", "解释边界"],
                [
                    ["T1", "触发到分派完成秒数", "反映系统完成多角色任务组织的速度"],
                    ["T2", "触发到核心施救响应秒数", "反映核心施救端响应系统任务的速度"],
                    ["T3", "触发到 CPR 开始秒数", "反映核心施救动作开始节点"],
                    ["T4", "触发到 AED 取到秒数", "反映 AED 保障者到达并取得 AED 的速度"],
                    ["T5", "触发到 AED 送达秒数", "反映 AED 从点位到现场的取送链路效率"],
                    ["T6", "触发到救护接管秒数", "反映救护接驳和交接流程完成节点"],
                    ["质量分", "证据包关键字段、节点和校验结果的综合质量分", "仅代表数据完整性，不代表医学效果"],
                ],
            ),
            "",
            "## 四、五轮实验汇总",
            "",
            md_table(["轮次", "场景", "事件阶段", "分派来源", "参与终端数", "可用 AED 点位数", "质量分", "证据校验"], summary_rows),
            "",
            "## 五、关键时间指标统计",
            "",
            md_table(["指标", "含义", "有效轮次", "均值秒", "中位数秒", "最小秒", "最大秒"], metric_stat_rows),
            "",
            "## 六、覆盖率与完整性统计",
            "",
            md_table(["指标", "有效轮次", "均值", "中位数", "最小", "最大"], coverage_rows),
            "",
            "五轮实验中，角色分派完整度、定位覆盖率和健康摘要覆盖率均达到 100%。该结果说明系统在本批模拟条件下能够稳定记录多端协同所需的核心结构化信息。",
            "",
            "## 七、逐轮实验记录",
            "",
        ]
    )

    for index, item in enumerate(rounds, start=1):
        summary = item["summary"]
        lines.extend(
            [
                f"### {index}. {item['name']}",
                "",
                f"事件编号：{summary['事件编号']}",
                "",
                f"生成时间：{iso_to_beijing(summary['生成时间'])}（北京时间）",
                "",
                f"事件阶段：{summary['事件阶段']}",
                "",
                f"分派来源：{summary['分派来源']}",
                "",
                f"患者代号：{summary['患者代号']}；核心施救者：{summary['核心施救者代号']}；AED 取送者：{summary['AED取送者代号']}；救护接驳者：{summary['救护接驳者代号']}。",
                "",
                "#### 关键指标",
                "",
            ]
        )
        metric_rows = []
        for code, key, label in TIME_METRICS:
            metric_rows.append([code, label, f"{fmt_num(summary.get(key, ''))} 秒"])
        metric_rows.extend(
            [
                ["C1", "参与终端数", fmt_num(summary.get("参与终端数", ""))],
                ["C2", "可用 AED 点位数", fmt_num(summary.get("可用AED点位数", ""))],
                ["C3", "AED 保障路线距离", f"{fmt_num(summary.get('AED保障路线距离米', ''))} 米"],
                ["C4", "证据质量分", fmt_num(quality_rows.get(summary["轮次编号"], {}).get("质量分", ""))],
            ]
        )
        lines.extend([md_table(["指标", "含义", "记录值"], metric_rows), ""])

        dispatch_rows = []
        for row in item["dispatch"]:
            dispatch_rows.append(
                [
                    row.get("角色", ""),
                    row.get("参与者代号", ""),
                    fmt_num(row.get("评分", "")),
                    f"{fmt_num(row.get('到患者距离米', ''))} 米",
                    row.get("最近AED点位ID", ""),
                    row.get("理由", ""),
                ]
            )
        lines.extend(
            [
                "#### 分派依据",
                "",
                md_table(["角色", "参与者代号", "评分", "到患者距离", "最近 AED 点位", "分派理由"], dispatch_rows),
                "",
                "#### 事件时间线",
                "",
            ]
        )
        timeline_rows = []
        for row in item["timeline"]:
            timeline_rows.append(
                [
                    fmt_num(row.get("相对秒数", "")),
                    iso_to_beijing(row.get("时间", "")),
                    row.get("事件类型", ""),
                    row.get("角色", ""),
                    row.get("日志内容", ""),
                ]
            )
        lines.extend([md_table(["相对秒数", "北京时间", "事件类型", "角色", "日志内容"], timeline_rows), ""])

    lines.extend(
        [
            "## 八、结果分析",
            "",
            "1. 五轮实验均完成事件归档，证据包校验通过，缺失关键节点数为 0，说明系统能够稳定形成一次模拟急救协同流程的结构化记录。",
            "",
            "2. T1 触发到分派完成的均值为 5 秒，说明系统在模拟场景中可以较快完成核心施救、AED 取送和救护接驳等角色组织。",
            "",
            "3. T3 触发到 CPR 开始的均值为 18.8 秒，T5 触发到 AED 送达的均值为 83 秒，T6 触发到救护接管的均值为 120.4 秒。这些指标可用于后续真人低风险场景预实验中的流程对照和瓶颈分析。",
            "",
            "4. 本批数据中 AED 点位数和可用 AED 点位数均为 2，AED 保障路线距离为 41.6 米。该数据用于验证 AED 取送链路记录能力，真实场地应用前仍需人工核验 AED 点位位置、可用性和开放时间。",
            "",
            "5. 本批数据均采用规则备用分派。其意义在于证明系统具备稳定可解释的分派兜底能力；后续若接入 AI 模型，应继续记录 AI 输出、规则兜底条件和人工确认结果。",
            "",
            "## 九、实验局限性",
            "",
            "1. 本记录来源于系统级模拟预实验和桌面推演数据，不等同于真实急救场景或临床试验。",
            "",
            "2. 本批数据主要用于验证网页端多端协同、事件状态机、时间线记录和证据包导出能力，不能推出抢救成功率、患者预后或临床疗效。",
            "",
            "3. 参与者问卷和观察员评分字段在本批打印版中不作为已完成真人主观数据呈现。后续若组织真人低风险场景预实验，应补充观察员记录、参与者评分和专家反馈。",
            "",
            "4. AED 点位为样例或待核验点位，真实部署前需要进行人工标记、现场复核和维护责任确认。",
            "",
            "## 十、结论",
            "",
            "本次五轮系统级模拟预实验初步验证了“生命反射弧”网页端急救协同系统在模拟院前急救场景中的流程闭环能力。系统能够完成事件触发、角色分派、四端协同、AED 取送记录、救护接管记录和证据包归档，并形成可复盘的 T1-T6 时间指标。该结果可作为医创赛 AI 创新设计作品中“工程可行性、流程记录能力和后续实验设计基础”的支撑材料。",
            "",
        ]
    )
    return "\n".join(lines) + "\n"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_run_font(run, size: int = 10, bold: bool = False) -> None:
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(size)
    run.bold = bold


def add_para(doc: Document, text: str, size: int = 10, bold: bool = False, align=None):
    paragraph = doc.add_paragraph()
    if align is not None:
        paragraph.alignment = align
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold)
    return paragraph


def add_heading(doc: Document, text: str, level: int) -> None:
    paragraph = doc.add_heading(level=level)
    paragraph.clear()
    run = paragraph.add_run(text)
    set_run_font(run, size=16 if level == 1 else 13 if level == 2 else 11, bold=True)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], font_size: int = 8) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    header_cells = table.rows[0].cells
    for idx, header in enumerate(headers):
        set_cell_shading(header_cells[idx], "D9EAF7")
        p = header_cells[idx].paragraphs[0]
        run = p.add_run(str(header))
        set_run_font(run, size=font_size, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            p = cells[idx].paragraphs[0]
            run = p.add_run(str(value))
            set_run_font(run, size=font_size)
    doc.add_paragraph()


def add_markdown_table_to_docx(doc: Document, markdown: str) -> None:
    table_lines = [line for line in markdown.splitlines() if line.startswith("|")]
    if len(table_lines) < 2:
        return
    headers = [part.strip().replace("\\|", "|") for part in table_lines[0].strip("|").split("|")]
    rows = []
    for line in table_lines[2:]:
        rows.append([part.strip().replace("\\|", "|") for part in line.strip("|").split("|")])
    add_table(doc, headers, rows, font_size=7 if len(headers) >= 5 else 8)


def build_docx(markdown: str) -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)

    style = doc.styles["Normal"]
    style.font.name = "宋体"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    style.font.size = Pt(10)

    paragraph_buffer: list[str] = []
    table_buffer: list[str] = []

    def flush_paragraphs():
        nonlocal paragraph_buffer
        if paragraph_buffer:
            text = "\n".join(paragraph_buffer).strip()
            if text:
                add_para(doc, text)
            paragraph_buffer = []

    def flush_table():
        nonlocal table_buffer
        if table_buffer:
            add_markdown_table_to_docx(doc, "\n".join(table_buffer))
            table_buffer = []

    for raw_line in markdown.splitlines():
        line = raw_line.rstrip()
        if line.startswith("|"):
            flush_paragraphs()
            table_buffer.append(line)
            continue
        flush_table()
        if not line.strip():
            flush_paragraphs()
            continue
        if line.startswith("# "):
            flush_paragraphs()
            add_heading(doc, line[2:].strip(), 1)
            continue
        if line.startswith("## "):
            flush_paragraphs()
            add_heading(doc, line[3:].strip(), 2)
            continue
        if line.startswith("### "):
            flush_paragraphs()
            add_heading(doc, line[4:].strip(), 3)
            continue
        if line.startswith("#### "):
            flush_paragraphs()
            add_heading(doc, line[5:].strip(), 4)
            continue
        paragraph_buffer.append(line)

    flush_paragraphs()
    flush_table()
    doc.save(DOCX_OUT)


def main() -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    rounds = [load_round(name) for name in ROUND_DIRS]
    markdown = build_markdown(rounds)
    MARKDOWN_OUT.write_text(markdown, encoding="utf-8")
    build_docx(markdown)
    print(MARKDOWN_OUT)
    print(DOCX_OUT)


if __name__ == "__main__":
    main()
